"""Create DOCX tool."""
from typing import Any, Dict
from .base import Tool, ToolResult


class CreateDocxTool(Tool):
    """Tool for creating Word documents."""

    def __init__(self):
        super().__init__(
            name="create_docx",
            description="Create a Word DOCX document. Input: JSON with 'content' (text), 'title' (optional), 'output' (filename).",
        )

    async def execute(self, **kwargs) -> ToolResult:
        try:
            from docx import Document
            
            content = kwargs.get("content", "")
            title = kwargs.get("title", "")
            output = kwargs.get("output", "document.docx")
            
            doc = Document()
            
            if title:
                doc.add_heading(title, 0)
            
            if isinstance(content, list):
                for line in content:
                    doc.add_paragraph(str(line))
            else:
                for line in str(content).split('\n'):
                    doc.add_paragraph(line)
            
            doc.save(output)
            return ToolResult(success=True, result=f"Created {output}")
        except ImportError:
            return ToolResult(success=False, result=None, error="python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            return ToolResult(success=False, result=None, error=str(e))

    def _get_parameters_schema(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "content": {"type": "string", "description": "Text content for document"},
                "title": {"type": "string", "description": "Document title (optional)"},
                "output": {"type": "string", "description": "Output filename (default: document.docx)"}
            },
            "required": ["content"]
        }